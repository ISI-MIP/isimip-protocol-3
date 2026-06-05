import sys
from pathlib import Path

from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from docx import Document

sys.path.insert(0, 'build')

from utils import filter_rows, read_definitions

SIMULATION_ROUNDS = ['ISIMIP4a', 'ISIMIP4b']
PRODUCT = 'OutputData'
IDENTIFIER = 'experiments'

COLOR_WHITE = 'FFFFFF'
COLOR_BLACK = '000000'
COLOR_HISTORICAL = 'e2e3e5'
COLOR_FUTURE = 'f8d7da'
COLOR_GROUP3 = 'cff4fc'

COLUMNS = {
    'ISIMIP4a': ['Experiment', 'Period: Historical'],
    'ISIMIP4b': ['Experiment', 'Period: Pre-industrial', 'Period: Historical', 'Period: Future'],
}

definitions = read_definitions()


def add_paragraph(cell, value, font_size=10, color=COLOR_BLACK, bold=False):
    value = value.replace('*', '')

    paragraph = cell.add_paragraph('') if cell.paragraphs[0].runs else cell.paragraphs[0]
    run = paragraph.add_run(value)

    run.font.name = 'Arial'
    run.font.size = Pt(font_size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_background(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)


def set_table_cell_margin(table, top=0, bottom=0, left=0, right=0):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)

    tblCellMar = tblPr.find(qn('w:tblCellMar'))
    if tblCellMar is None:
        tblCellMar = OxmlElement('w:tblCellMar')
        tblPr.append(tblCellMar)

    for attr, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        el = tblCellMar.find(qn(f'w:{attr}'))
        if el is None:
            el = OxmlElement(f'w:{attr}')
            tblCellMar.append(el)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')


for simulation_round in SIMULATION_ROUNDS:
    rows = list(filter_rows(definitions[IDENTIFIER], simulation_round, PRODUCT))
    columns = COLUMNS[simulation_round]

    doc = Document()
    table = doc.add_table(rows=len(rows) * 2 + 1, cols=len(columns))
    table.style = 'Table Grid'
    set_table_cell_margin(table, top=100, bottom=100, left=100, right=100)

    for row in table.rows:
        row.height = Cm(2)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    for column_index, column in enumerate(columns):
        cell = table.rows[0].cells[column_index]
        add_paragraph(cell, column, color=COLOR_WHITE, bold=True)
        set_background(cell, COLOR_BLACK)

    for row_index, row in enumerate(rows):
        index = (row_index * 2) + 1

        table.cell(index, 0).merge(table.cell(index + 1, 0))

        cell = table.rows[index].cells[0]
        add_paragraph(cell, row['title'], font_size=12, bold=True)

        cell = table.rows[index].cells[0]
        for subtitle in row['subtitles']:
            add_paragraph(cell, subtitle, bold=True)

        cell = table.rows[index].cells[0]
        add_paragraph(cell, row['priority'])

        if simulation_round == 'ISIMIP4a':
            for offset, key in enumerate(['climate', 'soc']):
                cell = table.rows[index + offset].cells[1]
                add_paragraph(cell, row['historical'][key], bold=True)
                set_background(cell, COLOR_HISTORICAL)

        else:
            if isinstance(row['pre-industrial'], str):
                table.cell(index, 1).merge(table.cell(index + 1, 1))
                cell = table.rows[index].cells[1]
                add_paragraph(cell, row['pre-industrial'])
            else:
                for offset, key in enumerate(['climate', 'soc']):
                    cell = table.rows[index + offset].cells[1]
                    add_paragraph(cell, row['pre-industrial'][key], bold=True)
                    if row['pre-industrial'].get(f'{key}_sens'):
                        value = 'Sensitivity experiment: ' + row['pre-industrial'][f'{key}_sens']
                        add_paragraph(cell, value, bold=True)
                    set_background(cell, COLOR_HISTORICAL)

            if isinstance(row['historical'], str):
                table.cell(index, 2).merge(table.cell(index + 1, 2))
                cell = table.rows[index].cells[2]
                add_paragraph(cell, row['historical'])
            else:
                for offset, key in enumerate(['climate', 'soc']):
                    cell = table.rows[index + offset].cells[2]
                    add_paragraph(cell, row['historical'][key], bold=True)
                    if row['historical'].get(f'{key}_sens'):
                        value = 'Sensitivity experiment: ' + row['historical'][f'{key}_sens']
                        add_paragraph(cell, value, bold=True)
                    set_background(cell, COLOR_HISTORICAL)

            if isinstance(row['future'], str):
                table.cell(index, 3).merge(table.cell(index + 1, 3))
                cell = table.rows[index].cells[3]
                add_paragraph(cell, row['future'])
            else:
                for offset, key in enumerate(['climate', 'soc']):
                    cell = table.rows[index + offset].cells[3]
                    add_paragraph(cell, row['future'][key], bold=True)
                    if row['future'].get(f'{key}_sens'):
                        value = 'Sensitivity experiment: ' + row['future'][f'{key}_sens']
                        add_paragraph(cell, value, bold=True)
                    set_background(cell, COLOR_GROUP3 if row['future']['soc'].endswith('adapt') else COLOR_FUTURE)

    output_path = Path('docx') / f'{simulation_round}-{IDENTIFIER}.docx'
    output_path.parent.mkdir(exist_ok=True)
    doc.save(output_path)
