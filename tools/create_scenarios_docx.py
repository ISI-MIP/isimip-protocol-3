import sys
from pathlib import Path

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from docx import Document

sys.path.insert(0, 'build')

from utils import filter_rows, read_definitions

SIMULATION_ROUNDS = ['ISIMIP4a', 'ISIMIP4b']
PRODUCT = 'OutputData'
IDENTIFIERS = ['climate_scenario', 'soc_scenario', 'sens_scenario']

COLOR_WHITE = 'FFFFFF'
COLOR_BLACK = '000000'

COLUMNS = ['Experiment specifier', 'Description']

definitions = read_definitions()


def add_paragraph(cell, value, font_size=10, color=COLOR_BLACK, bold=False, space_before=None):
    value = value.replace('**', '')

    paragraph = cell.add_paragraph('') if cell.paragraphs[0].runs else cell.paragraphs[0]
    run = paragraph.add_run(value)

    run.font.name = 'Calibri'
    run.font.size = Pt(font_size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)

    if space_before:
        paragraph.paragraph_format.space_before = Pt(space_before)


def set_background(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)


for simulation_round in SIMULATION_ROUNDS:
    for identifier in IDENTIFIERS:
        rows = list(filter_rows(definitions[identifier], simulation_round, PRODUCT))

        doc = Document()
        table = doc.add_table(rows=len(rows) + 1, cols=len(COLUMNS))
        table.style = 'Table Grid'

        for column_index, column in enumerate(COLUMNS):
            cell = table.rows[0].cells[column_index]
            add_paragraph(cell, column, color=COLOR_WHITE, bold=True)
            set_background(cell, COLOR_BLACK)

        for row_index, row in enumerate(rows):
            cell = table.rows[row_index + 1].cells[0]
            add_paragraph(cell, row['specifier'], font_size=12, bold=True)

            cell = table.rows[row_index + 1].cells[1]
            add_paragraph(cell, row['description'])

            if row.get('description_note'):
                add_paragraph(cell, row['description_note'], space_before=6)

        output_path = Path('docx') / f'{simulation_round}-{identifier}.docx'
        output_path.parent.mkdir(exist_ok=True)
        doc.save(output_path)
